"""
Document Processor Module
Handles text extraction from various document formats.
"""

import os
import logging
from typing import Optional, Dict, Any
import fitz  # PyMuPDF
from docx import Document
import markdown
from bs4 import BeautifulSoup
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing for multiple file formats."""
    
    ALLOWED_EXTENSIONS = {'adoc', 'md', 'dita', 'xml', 'docx', 'pdf', 'txt'}
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.md': self._extract_from_markdown,
            '.adoc': self._extract_from_asciidoc,
            '.dita': self._extract_from_dita,
            '.xml': self._extract_from_xml,
            '.txt': self._extract_from_text
        }
    
    def allowed_file(self, filename: str) -> bool:
        """Check if file extension is allowed."""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.ALLOWED_EXTENSIONS
    
    def extract_text(self, filepath: str) -> Optional[str]:
        """
        Extract text from a document file.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            if not os.path.exists(filepath):
                logger.error(f"File not found: {filepath}")
                return None
            
            file_ext = os.path.splitext(filepath)[1].lower()
            
            if file_ext not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_ext}")
                return None
            
            # Extract text using appropriate method
            text = self.supported_formats[file_ext](filepath)
            
            if text is not None:
                # Clean and normalize the text
                text = self._clean_text(text)
                logger.info(f"Successfully extracted {len(text)} characters from {filepath}")
                return text
            else:
                logger.error(f"Failed to extract text from {filepath}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, filepath: str) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(filepath)
            text = ""
            
            logger.info(f"🔍 [PDF-DEBUG] Extracting from PDF with {doc.page_count} pages")
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                try:
                    # Try newer PyMuPDF API first
                    text += page.get_text()
                except AttributeError:
                    # Fallback to older API
                    text += page.getText()
                text += "\n\n"  # Add page break
            
            doc.close()
            
            logger.info(f"🔍 [PDF-DEBUG] Extracted {len(text)} chars from PDF")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            return None
    
    def _extract_from_docx(self, filepath: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(filepath)
            text = ""
            
            logger.info(f"🔍 [DOCX-DEBUG] Extracting from DOCX with {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables")
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"🔍 [DOCX-DEBUG] Extracted {len(text)} chars from DOCX")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            return None
    
    def _extract_from_markdown(self, filepath: str) -> Optional[str]:
        """
        Extract text from Markdown file.
        CRITICAL: Preserves original Markdown markup for proper structural parsing.
        The Markdown parser (markdown/parser.py) will handle parsing and structure extraction.
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            logger.info(f"🔍 [MARKDOWN-DEBUG] Extracted raw Markdown content: {len(md_content)} chars")
            logger.info(f"🔍 [MARKDOWN-DEBUG] First 300 chars: {md_content[:300]}")
            
            # Return the content as-is, preserving all Markdown markup
            # This allows the structural parser to properly parse headings, lists, code blocks, etc.
            return md_content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from Markdown: {str(e)}")
            return None
    
    def _extract_from_asciidoc(self, filepath: str) -> Optional[str]:
        """
        Extract text from AsciiDoc file.
        CRITICAL: Preserves original AsciiDoc markup for proper structural parsing.
        The AsciiDoc parser (asciidoc/parser.py) will handle parsing and structure extraction.
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()
            
            logger.info(f"🔍 [ASCIIDOC-DEBUG] Extracted raw AsciiDoc content: {len(content)} chars")
            logger.info(f"🔍 [ASCIIDOC-DEBUG] First 300 chars: {content[:300]}")
            
            # Return the content as-is, preserving all AsciiDoc markup
            # This allows the structural parser to properly parse sections, tables, admonitions, etc.
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from AsciiDoc: {str(e)}")
            return None
    
    def _extract_from_dita(self, filepath: str) -> Optional[str]:
        """Extract text from DITA file."""
        return self._extract_from_xml(filepath)  # Use common XML extraction logic
    
    def _extract_from_xml(self, filepath: str) -> Optional[str]:
        """
        Extract text from XML file (including DITA).
        CRITICAL: Preserves original XML/DITA markup for proper structural parsing.
        The DITA parser (dita/parser.py) will handle parsing and structure extraction.
        """
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Check if this is DITA content
            is_dita = self._is_dita_content(content)
            
            if is_dita:
                logger.info(f"🔍 [DITA-DEBUG] Extracted raw DITA content: {len(content)} chars")
                logger.info(f"🔍 [DITA-DEBUG] First 300 chars: {content[:300]}")
            else:
                logger.info(f"🔍 [XML-DEBUG] Extracted raw XML content: {len(content)} chars")
                logger.info(f"🔍 [XML-DEBUG] First 300 chars: {content[:300]}")
            
            # Return the content as-is, preserving all XML/DITA markup
            # This allows the structural parser to properly parse topics, sections, elements, etc.
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from XML/DITA: {str(e)}")
            return None
    
    def _is_dita_content(self, content: str) -> bool:
        """Check if XML content is DITA format."""
        # Check for DITA DOCTYPE declarations
        dita_doctypes = [
            'DITA Concept', 'DITA Task', 'DITA Reference', 'DITA Topic',
            'DITA Troubleshooting', 'DITA Glossary', 'DITA Map'
        ]
        
        for doctype in dita_doctypes:
            if doctype in content:
                return True
        
        # Check for DITA root elements
        dita_roots = ['<concept', '<task', '<reference', '<topic', '<troubleshooting', '<map']
        content_lower = content.lower()
        
        for root in dita_roots:
            if root in content_lower:
                return True
        
        return False
    
    def _extract_from_text(self, filepath: str) -> Optional[str]:
        """Extract text from plain text file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()
            
            logger.info(f"🔍 [TEXT-DEBUG] Extracted plain text content: {len(content)} chars")
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from text file: {str(e)}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        NOTE: For structured formats (AsciiDoc, Markdown, DITA), minimal cleaning is applied
        to preserve formatting that structural parsers need.
        """
        if text is None:
            return ""
        if not text:
            return ""
        
        # Only trim leading/trailing whitespace
        # Do NOT collapse internal whitespace or newlines - structural parsers need them!
        text = text.strip()
        
        return text
    
    def extract_text_from_upload(self, file_storage) -> Optional[str]:
        """
        Extract text from an uploaded file (werkzeug FileStorage object).
        This method processes the file without saving it to disk.
        
        Args:
            file_storage: werkzeug FileStorage object from request.files
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            import tempfile
            
            # Get file extension from filename
            filename = file_storage.filename
            logger.info(f"🔍 [UPLOAD-DEBUG] Starting extraction for file: {filename}")
            
            if not filename or '.' not in filename:
                logger.error("❌ [UPLOAD-DEBUG] Invalid filename or no extension")
                return None
            
            file_ext = os.path.splitext(filename)[1].lower()
            logger.info(f"🔍 [UPLOAD-DEBUG] Detected file extension: {file_ext}")
            
            if file_ext not in self.supported_formats:
                logger.error(f"❌ [UPLOAD-DEBUG] Unsupported file format: {file_ext}")
                logger.error(f"❌ [UPLOAD-DEBUG] Supported formats: {list(self.supported_formats.keys())}")
                return None
            
            # Create temporary file for processing
            with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as tmp_file:
                tmp_path = tmp_file.name
                file_storage.save(tmp_path)
                logger.info(f"🔍 [UPLOAD-DEBUG] Saved to temporary file: {tmp_path}")
            
            try:
                # Extract text using the appropriate method
                logger.info(f"🔍 [UPLOAD-DEBUG] Calling extraction method for {file_ext}")
                text = self.supported_formats[file_ext](tmp_path)
                
                if text is not None:
                    logger.info(f"✅ [UPLOAD-DEBUG] Raw extraction successful: {len(text)} characters")
                    logger.info(f"🔍 [UPLOAD-DEBUG] First 200 chars of extracted text: {text[:200]}")
                    
                    # Clean and normalize the text
                    text = self._clean_text(text)
                    logger.info(f"✅ [UPLOAD-DEBUG] After cleaning: {len(text)} characters")
                    logger.info(f"🔍 [UPLOAD-DEBUG] First 200 chars after cleaning: {text[:200]}")
                    
                    return text
                else:
                    logger.error(f"❌ [UPLOAD-DEBUG] Extraction method returned None for {filename}")
                    return None
            
            finally:
                # Clean up temporary file
                try:
                    os.unlink(tmp_path)
                    logger.info(f"🔍 [UPLOAD-DEBUG] Cleaned up temporary file")
                except Exception as e:
                    logger.warning(f"⚠️ [UPLOAD-DEBUG] Failed to delete temporary file: {e}")
                    
        except Exception as e:
            logger.error(f"❌ [UPLOAD-DEBUG] Error extracting text from upload: {str(e)}")
            import traceback
            logger.error(f"❌ [UPLOAD-DEBUG] Traceback: {traceback.format_exc()}")
            return None
    
    def detect_format_from_filename(self, filename: str) -> str:
        """
        Detect document format from filename.
        
        Args:
            filename: Name of the file
            
        Returns:
            Format string compatible with parser_factory ('asciidoc', 'markdown', 'dita', 'plaintext', 'auto')
        """
        logger.info(f"🔍 [FORMAT-DEBUG] Detecting format for filename: {filename}")
        
        if not filename or '.' not in filename:
            logger.warning(f"⚠️ [FORMAT-DEBUG] No extension found, returning 'auto'")
            return 'auto'
        
        file_ext = os.path.splitext(filename)[1].lower()
        logger.info(f"🔍 [FORMAT-DEBUG] File extension: {file_ext}")
        
        # Map file extensions to parser format hints
        format_mapping = {
            '.adoc': 'asciidoc',
            '.asciidoc': 'asciidoc',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.dita': 'dita',
            '.xml': 'dita',  # Assume XML is DITA unless content says otherwise
            '.txt': 'plaintext',
            '.pdf': 'auto',  # PDF requires content analysis
            '.docx': 'auto'  # DOCX requires content analysis
        }
        
        detected_format = format_mapping.get(file_ext, 'auto')
        logger.info(f"✅ [FORMAT-DEBUG] Mapped {file_ext} → {detected_format}")
        
        return detected_format
    
    def get_document_info(self, filepath: str) -> Dict[str, Any]:
        """
        Get document information and metadata.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Dictionary with document information
        """
        info = {
            'filepath': filepath,
            'filename': os.path.basename(filepath),
            'file_size': 0,
            'format': 'unknown',
            'extractable': False
        }
        
        try:
            # Extract format from file extension regardless of file existence
            file_ext = os.path.splitext(filepath)[1].lower()
            if file_ext:
                info['format'] = file_ext.lstrip('.')
                info['extractable'] = file_ext in self.supported_formats
            
            if os.path.exists(filepath):
                info['file_size'] = os.path.getsize(filepath)
            
        except Exception as e:
            logger.error(f"Error getting document info: {str(e)}")
        
        return info 